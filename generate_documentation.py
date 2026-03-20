from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_with_color(doc, text, level=1, color=(0, 102, 204)):
    """Add a colored heading to the document"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_code_block(doc, code, language=""):
    """Add a code block with gray background"""
    paragraph = doc.add_paragraph()
    paragraph.style = 'Normal'
    
    # Add language label if provided
    if language:
        run = paragraph.add_run(f"{language}\n")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
        run.italic = True
    
    # Add code content
    run = paragraph.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    # Set paragraph shading (background color)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F0F0F0')
    paragraph._element.get_or_add_pPr().append(shading_elm)
    
    return paragraph

def create_project_documentation():
    """Create comprehensive Word document for the Spring Boot Chatbot project"""
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('Spring Boot Chatbot Application', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.size = Pt(28)
    
    subtitle = doc.add_paragraph('Web-Based AI Assistant with Gemini API Integration')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project metadata
    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_text = """
Project Type: Full-Stack Web Application
Technology Stack: Spring Boot, Java 17, Maven, Docker
API Integration: Google Gemini Pro
Frontend: HTML5, CSS3, JavaScript
Deployment: Docker Containerization
    """
    metadata.add_run(metadata_text).font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== TABLE OF CONTENTS =====
    add_heading_with_color(doc, 'Table of Contents', 1, (0, 102, 204))
    toc_items = [
        "1. Project Overview",
        "2. Technology Stack",
        "3. Project Architecture",
        "4. Key Features",
        "5. Project Structure",
        "6. Source Code Explanation",
        "   6.1 Main Application Class",
        "   6.2 Controller Layer",
        "   6.3 Service Layer",
        "   6.4 Model Classes",
        "   6.5 Frontend (HTML/CSS/JavaScript)",
        "7. Docker Containerization",
        "8. Application Screenshots",
        "9. Setup and Installation",
        "10. Testing the Application",
        "11. Conclusion"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 1. PROJECT OVERVIEW =====
    add_heading_with_color(doc, '1. Project Overview', 1, (0, 102, 204))
    
    doc.add_paragraph(
        "This project is a modern, full-stack web-based chatbot application built using Spring Boot framework. "
        "It demonstrates the integration of a backend REST API with Google's Gemini Pro AI model to create an "
        "intelligent conversational assistant."
    )
    
    doc.add_paragraph(
        "The application showcases several important software development concepts:"
    )
    
    features_list = [
        "RESTful API design and implementation",
        "Integration with third-party APIs (Google Gemini)",
        "Layered architecture (Controller → Service → Model)",
        "Responsive web interface with modern UI/UX",
        "Docker containerization for easy deployment",
        "Maven-based project management",
        "Hardcoded fallback responses for offline testing"
    ]
    
    for feature in features_list:
        doc.add_paragraph(feature, style='List Bullet')
    
    # ===== 2. TECHNOLOGY STACK =====
    add_heading_with_color(doc, '2. Technology Stack', 1, (0, 102, 204))
    
    add_heading_with_color(doc, 'Backend Technologies', 2, (0, 153, 76))
    backend_tech = [
        "Java 17 - Programming language",
        "Spring Boot 3.2.0 - Application framework",
        "Spring Web - RESTful web services",
        "Spring WebFlux - Reactive HTTP client for API calls",
        "Thymeleaf - Server-side template engine",
        "Maven - Build automation and dependency management"
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    add_heading_with_color(doc, 'Frontend Technologies', 2, (0, 153, 76))
    frontend_tech = [
        "HTML5 - Structure and markup",
        "CSS3 - Styling with modern design patterns",
        "JavaScript (ES6+) - Client-side interactivity",
        "Font Awesome - Icon library",
        "Google Fonts (Outfit) - Typography"
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    add_heading_with_color(doc, 'DevOps & Deployment', 2, (0, 153, 76))
    devops_tech = [
        "Docker - Containerization",
        "Docker Compose - Multi-container orchestration",
        "Multi-stage Docker builds - Optimized image size",
        "Health checks - Container monitoring"
    ]
    for tech in devops_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 3. PROJECT ARCHITECTURE =====
    add_heading_with_color(doc, '3. Project Architecture', 1, (0, 102, 204))
    
    doc.add_paragraph(
        "The application follows a layered architecture pattern, separating concerns into distinct layers:"
    )
    
    add_heading_with_color(doc, 'Architecture Layers', 2, (0, 153, 76))
    
    architecture_desc = """
1. Presentation Layer (Frontend)
   - HTML/CSS/JavaScript interface
   - Handles user interactions
   - Sends AJAX requests to backend API
   - Displays chat messages dynamically

2. Controller Layer
   - ChatController.java
   - Handles HTTP requests (GET, POST)
   - Routes requests to appropriate services
   - Returns JSON responses

3. Service Layer
   - ChatService.java
   - Contains business logic
   - Manages API integration with Gemini
   - Provides hardcoded responses for testing
   - Handles error scenarios

4. Model Layer
   - ChatRequest.java - Request data structure
   - ChatResponse.java - Response data structure
   - POJOs (Plain Old Java Objects)

5. External Integration
   - Google Gemini Pro API
   - WebClient for reactive HTTP calls
   - JSON request/response handling
    """
    doc.add_paragraph(architecture_desc)
    
    doc.add_page_break()
    
    # ===== 4. KEY FEATURES =====
    add_heading_with_color(doc, '4. Key Features', 1, (0, 102, 204))
    
    features = {
        "🤖 AI-Powered Responses": "Integration with Google Gemini Pro for intelligent conversations",
        "💬 Real-time Chat Interface": "Modern, responsive chat UI with message bubbles and avatars",
        "🎨 Theme Toggle": "Dark/Light mode support with persistent preferences",
        "📝 Hardcoded Responses": "Fallback responses for common queries (DevOps, Maven, Spring Boot)",
        "🔄 Loading Indicators": "Visual feedback during API calls with typing animation",
        "🗑️ Clear Chat": "Ability to clear conversation history",
        "⌨️ Keyboard Shortcuts": "Ctrl+K (focus input), Ctrl+L (clear chat), Ctrl+Shift+T (toggle theme)",
        "🐳 Docker Support": "Fully containerized with multi-stage builds",
        "❤️ Health Checks": "Container health monitoring for production deployments",
        "📱 Responsive Design": "Works seamlessly on desktop and mobile devices"
    }
    
    for feature_name, feature_desc in features.items():
        p = doc.add_paragraph()
        p.add_run(f"{feature_name}: ").bold = True
        p.add_run(feature_desc)
    
    doc.add_page_break()
    
    # ===== 5. PROJECT STRUCTURE =====
    add_heading_with_color(doc, '5. Project Structure', 1, (0, 102, 204))
    
    project_structure = """
chatbot/
├── src/
│   ├── main/
│   │   ├── java/com/example/chatbot/
│   │   │   ├── ChatbotApplication.java      # Main Spring Boot application
│   │   │   ├── controller/
│   │   │   │   └── ChatController.java      # REST API endpoints
│   │   │   ├── service/
│   │   │   │   └── ChatService.java         # Business logic & API integration
│   │   │   └── model/
│   │   │       ├── ChatRequest.java         # Request model
│   │   │       └── ChatResponse.java        # Response model
│   │   └── resources/
│   │       ├── application.properties       # Configuration
│   │       ├── templates/
│   │       │   └── index.html              # Main HTML page
│   │       └── static/
│   │           └── style.css               # Stylesheet
│   └── test/                               # Unit tests
├── pom.xml                                 # Maven dependencies
├── Dockerfile                              # Docker build instructions
├── docker-compose.yml                      # Container orchestration
└── DOCKER_FLOW.md                         # Docker documentation
    """
    add_code_block(doc, project_structure, "Project Directory Structure")
    
    doc.add_page_break()
    
    # ===== 6. SOURCE CODE EXPLANATION =====
    add_heading_with_color(doc, '6. Source Code Explanation', 1, (0, 102, 204))
    
    # 6.1 Main Application
    add_heading_with_color(doc, '6.1 Main Application Class', 2, (0, 153, 76))
    doc.add_paragraph(
        "The ChatbotApplication.java is the entry point of the Spring Boot application. "
        "It uses the @SpringBootApplication annotation which combines @Configuration, "
        "@EnableAutoConfiguration, and @ComponentScan."
    )
    
    main_app_code = """package com.example.chatbot;

import org.springframework.boot.SpringApplication;
import org.springframework.boot.autoconfigure.SpringBootApplication;

@SpringBootApplication
public class ChatbotApplication {
    public static void main(String[] args) {
        SpringApplication.run(ChatbotApplication.class, args);
    }
}"""
    add_code_block(doc, main_app_code, "ChatbotApplication.java")
    
    # 6.2 Controller Layer
    add_heading_with_color(doc, '6.2 Controller Layer', 2, (0, 153, 76))
    doc.add_paragraph(
        "The ChatController handles HTTP requests. It has two main endpoints:"
    )
    doc.add_paragraph("• GET / - Serves the main HTML page", style='List Bullet')
    doc.add_paragraph("• POST /api/chat - Handles chat messages and returns JSON responses", style='List Bullet')
    
    controller_code = """@Controller
public class ChatController {
    private final ChatService chatService;

    @Autowired
    public ChatController(ChatService chatService) {
        this.chatService = chatService;
    }

    @GetMapping("/")
    public String index() {
        return "index";  // Returns index.html template
    }

    @PostMapping("/api/chat")
    @ResponseBody
    public ResponseEntity<ChatResponse> chat(@RequestBody ChatRequest request) {
        if (request == null || request.getMessage() == null || 
            request.getMessage().trim().isEmpty()) {
            return ResponseEntity.badRequest()
                .body(new ChatResponse("Please provide a valid message."));
        }

        ChatResponse response = chatService.getChatResponse(request.getMessage());
        return ResponseEntity.ok(response);
    }
}"""
    add_code_block(doc, controller_code, "ChatController.java (Key Methods)")
    
    doc.add_page_break()
    
    # 6.3 Service Layer
    add_heading_with_color(doc, '6.3 Service Layer', 2, (0, 153, 76))
    doc.add_paragraph(
        "The ChatService contains the core business logic. It first checks for hardcoded responses "
        "for common queries (like 'hello', 'what is devops?'), and if no match is found, it calls "
        "the Gemini API using Spring WebFlux's WebClient."
    )
    
    service_code = """@Service
public class ChatService {
    @Value("${gemini.api.key}")
    private String apiKey;

    @Value("${gemini.api.url}")
    private String apiUrl;

    private final WebClient webClient;

    public ChatResponse getChatResponse(String message) {
        try {
            // Check for hardcoded responses first
            String response = getHardcodedResponse(message);
            if (response != null) {
                return new ChatResponse(response);
            }

            // Prepare Gemini API request
            Map<String, Object> requestBody = new HashMap<>();
            // ... build request structure ...

            // Make API call
            String fullUrl = apiUrl + "?key=" + apiKey;
            Map<String, Object> apiResponse = webClient.post()
                .uri(fullUrl)
                .bodyValue(requestBody)
                .retrieve()
                .bodyToMono(Map.class)
                .block();

            String responseText = extractResponseText(apiResponse);
            return new ChatResponse(responseText);

        } catch (Exception e) {
            return new ChatResponse("Error: " + e.getMessage());
        }
    }

    private String getHardcodedResponse(String message) {
        String lowerMessage = message.toLowerCase().trim();
        
        if (lowerMessage.equals("hello") || lowerMessage.equals("hi")) {
            return "Hello! 👋 Welcome to the Web-Based Chatbot.";
        }
        
        if (lowerMessage.contains("what is devops")) {
            return "DevOps is a set of practices that combines software " +
                   "development (Dev) and IT operations (Ops)...";
        }
        
        // ... more hardcoded responses ...
        return null;  // No match found
    }
}"""
    add_code_block(doc, service_code, "ChatService.java (Simplified)")
    
    doc.add_page_break()
    
    # 6.4 Model Classes
    add_heading_with_color(doc, '6.4 Model Classes', 2, (0, 153, 76))
    doc.add_paragraph("Simple POJOs (Plain Old Java Objects) for request and response data:")
    
    model_code = """// ChatRequest.java
public class ChatRequest {
    private String message;

    public String getMessage() { return message; }
    public void setMessage(String message) { this.message = message; }
}

// ChatResponse.java
public class ChatResponse {
    private String response;

    public ChatResponse(String response) {
        this.response = response;
    }

    public String getResponse() { return response; }
    public void setResponse(String response) { this.response = response; }
}"""
    add_code_block(doc, model_code, "Model Classes")
    
    # 6.5 Frontend
    add_heading_with_color(doc, '6.5 Frontend (HTML/CSS/JavaScript)', 2, (0, 153, 76))
    doc.add_paragraph(
        "The frontend consists of a modern, responsive chat interface with:"
    )
    frontend_features = [
        "Sidebar navigation with logo and user profile",
        "Chat message area with scrollable message history",
        "Input field with send button",
        "Theme toggle (dark/light mode)",
        "Clear chat functionality",
        "Loading animations and typing indicators",
        "Keyboard shortcuts for better UX"
    ]
    for feature in frontend_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    frontend_code = """// JavaScript - Sending a chat message
chatForm.addEventListener('submit', async (e) => {
    e.preventDefault();
    
    const message = messageInput.value.trim();
    if (!message) return;

    // Add user message to UI
    addMessage(message, 'user');
    messageInput.value = '';
    
    // Show loading indicator
    const loadingId = addLoadingBubble();

    try {
        // Send POST request to backend
        const response = await fetch('/api/chat', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ message: message })
        });

        const data = await response.json();
        
        // Remove loading and show bot response
        document.getElementById(loadingId)?.remove();
        addMessage(data.response, 'bot');
    } catch (error) {
        document.getElementById(loadingId)?.remove();
        addMessage('Sorry, something went wrong.', 'bot', true);
    }
});"""
    add_code_block(doc, frontend_code, "JavaScript - Chat Functionality")
    
    doc.add_page_break()
    
    # ===== 7. DOCKER CONTAINERIZATION =====
    add_heading_with_color(doc, '7. Docker Containerization', 1, (0, 102, 204))
    
    doc.add_paragraph(
        "The application uses a multi-stage Docker build to create an optimized production image:"
    )
    
    docker_stages = """
Stage 1: Builder
- Base image: maven:3.9.11-eclipse-temurin-17
- Copies source code and pom.xml
- Runs 'mvn clean package' to build JAR file
- Output: chatbot-1.0.0.jar (~45MB)

Stage 2: Runtime
- Base image: eclipse-temurin:17-jre-alpine (minimal JRE)
- Copies only the JAR file from builder stage
- Exposes port 8085
- Configures health checks
- Final image size: ~380MB (optimized)
    """
    doc.add_paragraph(docker_stages)
    
    dockerfile_code = """# Stage 1: Build stage
FROM maven:3.9.11-eclipse-temurin-17 AS builder
WORKDIR /build
COPY pom.xml .
COPY src ./src
RUN mvn clean package -DskipTests -q

# Stage 2: Runtime stage
FROM eclipse-temurin:17-jre-alpine
WORKDIR /app
COPY --from=builder /build/target/chatbot-1.0.0.jar app.jar
EXPOSE 8085

HEALTHCHECK --interval=30s --timeout=3s --start-period=5s --retries=3 \\
    CMD wget --no-verbose --tries=1 --spider http://localhost:8085/ || exit 1

ENTRYPOINT ["java", "-jar", "app.jar"]"""
    add_code_block(doc, dockerfile_code, "Dockerfile")
    
    doc.add_paragraph()
    doc.add_paragraph("Benefits of Docker containerization:")
    docker_benefits = [
        "Consistent environment across development and production",
        "Easy deployment and scaling",
        "Isolated dependencies",
        "Reproducible builds",
        "Smaller image size with multi-stage builds",
        "Built-in health monitoring"
    ]
    for benefit in docker_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 8. APPLICATION SCREENSHOTS =====
    add_heading_with_color(doc, '8. Application Screenshots', 1, (0, 102, 204))
    
    doc.add_paragraph(
        "Below are screenshots demonstrating the application in action:"
    )
    
    # Screenshot 1: Initial Page
    add_heading_with_color(doc, '8.1 Initial Landing Page', 2, (0, 153, 76))
    doc.add_paragraph(
        "The chatbot interface when first loaded, showing the welcome message and modern UI design:"
    )
    
    screenshot1_path = r"C:\Users\perfe\.gemini\antigravity\brain\cf9f5b8e-ded6-4744-93b9-01c32dbf4741\initial_page_1765960804340.png"
    if os.path.exists(screenshot1_path):
        try:
            doc.add_picture(screenshot1_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Screenshot could not be embedded: {e}]")
    
    doc.add_page_break()
    
    # Screenshot 2: Hello Response
    add_heading_with_color(doc, '8.2 Greeting Interaction', 2, (0, 153, 76))
    doc.add_paragraph(
        "User sends 'hello' and receives a friendly hardcoded response from the chatbot:"
    )
    
    screenshot2_path = r"C:\Users\perfe\.gemini\antigravity\brain\cf9f5b8e-ded6-4744-93b9-01c32dbf4741\hello_response_1765960825723.png"
    if os.path.exists(screenshot2_path):
        try:
            doc.add_picture(screenshot2_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Screenshot could not be embedded: {e}]")
    
    doc.add_page_break()
    
    # Screenshot 3: DevOps Question
    add_heading_with_color(doc, '8.3 Technical Question Response', 2, (0, 153, 76))
    doc.add_paragraph(
        "User asks 'what is devops?' and receives a detailed, informative response:"
    )
    
    screenshot3_path = r"C:\Users\perfe\.gemini\antigravity\brain\cf9f5b8e-ded6-4744-93b9-01c32dbf4741\devops_response_1765960846652.png"
    if os.path.exists(screenshot3_path):
        try:
            doc.add_picture(screenshot3_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Screenshot could not be embedded: {e}]")
    
    # Screenshot 4: Existing screenshot from project
    doc.add_page_break()
    add_heading_with_color(doc, '8.4 Additional Application View', 2, (0, 153, 76))
    doc.add_paragraph(
        "Another view of the application interface:"
    )
    
    screenshot4_path = r"c:\Users\perfe\Desktop\devops\Screenshot 2025-12-15 124938.png"
    if os.path.exists(screenshot4_path):
        try:
            doc.add_picture(screenshot4_path, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"[Screenshot could not be embedded: {e}]")
    
    doc.add_page_break()
    
    # ===== 9. SETUP AND INSTALLATION =====
    add_heading_with_color(doc, '9. Setup and Installation', 1, (0, 102, 204))
    
    add_heading_with_color(doc, '9.1 Prerequisites', 2, (0, 153, 76))
    prerequisites = [
        "Java 17 or higher",
        "Maven 3.6+",
        "Docker Desktop (for containerized deployment)",
        "Google Gemini API key (optional, for live API integration)"
    ]
    for prereq in prerequisites:
        doc.add_paragraph(prereq, style='List Bullet')
    
    add_heading_with_color(doc, '9.2 Running with Maven', 2, (0, 153, 76))
    maven_steps = """
1. Clone or download the project
2. Navigate to project directory
3. Configure API key in src/main/resources/application.properties:
   gemini.api.key=your_api_key_here
4. Run the application:
   mvn spring-boot:run
5. Open browser and navigate to:
   http://localhost:8085
    """
    doc.add_paragraph(maven_steps)
    
    add_heading_with_color(doc, '9.3 Running with Docker', 2, (0, 153, 76))
    docker_steps = """
1. Build the Docker image:
   docker build -t chatbot:latest .

2. Run the container:
   docker run -d -p 8085:8085 chatbot:latest

3. Or use Docker Compose:
   docker-compose up --build

4. Access the application:
   http://localhost:8085
    """
    doc.add_paragraph(docker_steps)
    
    doc.add_page_break()
    
    # ===== 10. TESTING THE APPLICATION =====
    add_heading_with_color(doc, '10. Testing the Application', 1, (0, 102, 204))
    
    doc.add_paragraph(
        "The application includes hardcoded responses for testing without a valid API key. "
        "Try these test queries:"
    )
    
    test_queries = [
        "hello / hi / hey - Greeting responses",
        "what is devops? - Detailed DevOps explanation",
        "what is maven? - Maven build tool information",
        "what is spring boot? - Spring Boot framework details",
        "how are you? - Friendly response",
        "help - List of available topics",
        "thank you - Acknowledgment response"
    ]
    
    for query in test_queries:
        doc.add_paragraph(query, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        "For queries not in the hardcoded list, the application will attempt to call the "
        "Gemini API (if a valid API key is configured) or return a helpful error message."
    )
    
    doc.add_page_break()
    
    # ===== 11. CONCLUSION =====
    add_heading_with_color(doc, '11. Conclusion', 1, (0, 102, 204))
    
    conclusion_text = """
This Spring Boot Chatbot application demonstrates a complete full-stack development workflow, 
from backend API design to frontend user interface, with modern DevOps practices including 
Docker containerization.

Key Learning Outcomes:
• Understanding of Spring Boot framework and layered architecture
• REST API design and implementation
• Integration with third-party APIs (Google Gemini)
• Modern web UI development with responsive design
• Docker containerization and deployment strategies
• Maven project structure and dependency management

The application serves as an excellent foundation for building more complex chatbot systems 
and can be extended with features like:
• User authentication and session management
• Chat history persistence with database integration
• File upload and processing capabilities
• Multi-language support
• Advanced AI features and custom training
• Real-time notifications with WebSockets

This project showcases industry-standard practices and provides a solid understanding of 
modern web application development and deployment.
    """
    doc.add_paragraph(conclusion_text)
    
    # Add footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('─' * 80)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph('Spring Boot Chatbot Application Documentation')
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_text.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(100, 100, 100)
    footer_run.italic = True
    
    # Save document
    output_path = r"c:\Users\perfe\Desktop\devops\Spring_Boot_Chatbot_Documentation.docx"
    doc.save(output_path)
    print(f"Documentation created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_project_documentation()
